import docx

# 打开Word文档
doc = docx.Document("GPU显卡安装文档.docx")

# 读取文档中的段落
for paragraph in doc.paragraphs:
    print(f"paragraph: {paragraph.text}")

# 读取文档中的表格
# for table in doc.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             print(cell.text)
